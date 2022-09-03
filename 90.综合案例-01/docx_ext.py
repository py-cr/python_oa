# -*- coding: utf-8 -*-
# @Time    : 2022/4/23 15:30
# @Author  : Python超人
# @FileName: docx_ext.py

import docx # 导入word工具包
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX


def add_bookmark(paragraph, text, name):
    """
    在指定的段落对象后面增加书签
    例子：
    p1 = document.add_paragraph()
    hyperlink = add_bookmark(paragraph=p1,
                          text='书签文本',
                          name='书签名')
    :param paragraph: 段落对象
    :param text: 书签文本
    :param name: 书签名
    :return:
    """
    # tag = run._r  # for reference the following also works: tag =  document.element.xpath('//w:r')[-1]
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), name)
    start.set(docx.oxml.ns.qn('w:name'), name)
    paragraph._p.append(start)

    txt = docx.oxml.OxmlElement('w:r')
    txt.text = text
    paragraph._p.append(txt)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), name)
    end.set(docx.oxml.ns.qn('w:name'), name)

    paragraph._p.append(end)

def add_hyperlink(paragraph, url, text, tooltip=None, font=None, color=None, underline=True):
    """
    在指定的段落对象后面增加链接
    例子：
    p1 = document.add_paragraph()
    hyperlink = add_hyperlink(paragraph=p1,
                          url='http://www.baidu.com',
                          text='百度',
                          tooltip=None,
                          font="微软雅黑",
                          color=None,
                          underline=True)
    :param paragraph: 段落对象
    :param url: 链接地址
    :param text: 链接文本
    :param tooltip: 链接提示（为空则为 url）
    :param font: 字体
    :param color: 颜色
    :param underline: 下划线
    :return: 超链接对象
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), url, )
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    if tooltip:
        hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tooltip, )
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    # <w:color w:val="000000" w:themeColor="hyperlink" />
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)
    else:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:themeColor'), "hyperlink")
        rPr.append(c)

    if font:
        # <w:rFonts w:ascii="Arial" w:eastAsia="宋体" w:hAnsi="Arial" w:cs="Arial" />
        f = docx.oxml.shared.OxmlElement('w:rFonts')
        f.set(docx.oxml.shared.qn('w:ascii'), font)
        f.set(docx.oxml.shared.qn('w:eastAsia'), font)
        f.set(docx.oxml.shared.qn('w:hAnsi'), font)
        f.set(docx.oxml.shared.qn('w:cs'), font)
        rPr.append(f)

    if underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    new_run.bold = True
    # new_run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


#读取docx中的文本代码示例
def read_docx(docx_file):
    # 打开 Word 文档
    document = docx.Document(docx_file)
    # 读取 Word 文档中的段落
    # 文档中的 paragraphs 段落对象列表
    for paragraph in document.paragraphs:
        # 段落对象中的 runs 文本块对象列表
        for r in paragraph.runs:
            r_info = "" # 存放一个文本块内容和样式信息，print 结果
            if r.font.name is not None: # 如果字体名称不为空，显示字体名
                r_info += r.font.name + ","
            if r.font.size is not None: # 如果字号不为空，显示字号
                r_info += "字号:" + str(round(int(r.font.size)/12700)) + ","
            if r.bold: # 如果是粗体字
                r_info += "粗体字,"
            if r.italic: # 如果是斜体字
                r_info += "斜体字,"
            if r.underline: # 如果是下划线字
                r_info += "下划线,"
            if r.font.color.rgb is not None: # 如果文字有颜色，显示字体颜色
                r_info += "颜色:" + str(r.font.color.rgb) + ","
            # 如果长度大于0，说明 r_info 有样式信息
            if len(r_info) > 0:
                # r_info 中最后一个字符肯定有一个“逗号”
                r_info = r_info[0:-1] # 去掉末尾的“逗号”
                r_info = r.text + f"【{r_info}】"
            else: # 没有样式信息就直接显示文字
                r_info = r.text
            print(r_info)
            
    # 文档中的 tables 表格对象列表
    for table in document.tables:
        # 打印便于简单的展示表格
        print("---------------------------------")
        for row in table.rows:
            row_text = ""
            for cell in row.cells:
                row_text += cell.text + "\t|\t" # \t 代表的意思是水平制表符，能否保证列尽可能的对齐
            print(row_text[0:-3]) # [0:-3] 去掉最后的 \t|\t
            # 打印便于简单的展示表格
            print("---------------------------------")

    def show_img(blob):
        """
        显示图片的函数（以后的课程再讲解）
        """
        from PIL import Image # 图片处理工具包
        from io import BytesIO # 字节流工具包
        import matplotlib.pyplot as plt # 图片显示
        plt.figure('image',figsize=(8,6))
        plt.xticks([])
        plt.yticks([])
        im = Image.open(BytesIO(blob))
        plt.imshow(im)
        plt.show()

    # 读取 Word 文档中的图片（工具包中没有直接提供图片对象列表）
    dict_rel = document.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]

        if "image" in rel.reltype:
            show_img(rel.target_part.blob)
        else:
            # word文档非常复杂，暂不扩展课程内容
            # 我们把 word的扩展名改一下，您会发现确实复杂
            pass # 这个语法就是 pass 通过，什么事情都不做，以后会经常见到  

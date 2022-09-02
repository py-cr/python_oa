# -*- coding: utf-8 -*-
# @Time    : 2022/4/23 15:30
# @Author  : Python超人
# @FileName: docx_ext.py

import docx # 导入word工具包
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


#读取docx中的文本代码示例
def read_docx(docx_file):
    # 打开 Word 文档
    document = docx.Document(docx_file)
    # 读取 Word 文档中的段落
    # 文档中的 paragraphs 段落对象列表
    for paragraph in document.paragraphs:
        # 段落对象中的 runs 文本块对象列表
        for r in paragraph.runs:
            r_info = "" # 存放一个文本块内容和样式信息，print 结果
            if r.font.name is not None: # 如果字体名称不为空，显示字体名
                r_info += r.font.name + ","
            if r.font.size is not None: # 如果字号不为空，显示字号
                r_info += "字号:" + str(round(int(r.font.size)/12700)) + ","
            if r.bold: # 如果是粗体字
                r_info += "粗体字,"
            if r.italic: # 如果是斜体字
                r_info += "斜体字,"
            if r.font.color.rgb is not None: # 如果文字有颜色，显示字体颜色
                r_info += "颜色:" + str(r.font.color.rgb) + ","
            # 如果长度大于0，说明 r_info 有样式信息
            if len(r_info) > 0:
                # r_info 中最后一个字符肯定有一个“逗号”
                r_info = r_info[0:-1] # 去掉末尾的“逗号”
                r_info = r.text + f"【{r_info}】"
            else: # 没有样式信息就直接显示文字
                r_info = r.text
            print(r_info)
            
    # 文档中的 tables 表格对象列表
    for table in document.tables:
        # 打印便于简单的展示表格
        print("---------------------------------")
        for row in table.rows:
            row_text = ""
            for cell in row.cells:
                row_text += cell.text + "\t|\t" # \t 代表的意思是水平制表符，能否保证列尽可能的对齐
            print(row_text[0:-3]) # [0:-3] 去掉最后的 \t|\t
            # 打印便于简单的展示表格
            print("---------------------------------")

    def show_img(blob):
        """
        显示图片的函数（以后的课程再讲解）
        """
        from PIL import Image # 图片处理工具包
        from io import BytesIO # 字节流工具包
        import matplotlib.pyplot as plt # 图片显示
        plt.figure('image',figsize=(8,6))
        plt.xticks([])
        plt.yticks([])
        im = Image.open(BytesIO(blob))
        plt.imshow(im)
        plt.show()

    # 读取 Word 文档中的图片（工具包中没有直接提供图片对象列表）
    dict_rel = document.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]

        if "image" in rel.reltype:
            show_img(rel.target_part.blob)
        else:
            # word文档非常复杂，暂不扩展课程内容
            # 我们把 word的扩展名改一下，您会发现确实复杂
            pass # 这个语法就是 pass 通过，什么事情都不做，以后会经常见到  
